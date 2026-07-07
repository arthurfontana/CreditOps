"""Edição rica da demanda: WYSIWYG, permissões, histórico, anexos e export docx."""

from __future__ import annotations

import io

import pytest

from app.services import attachment_service, change_request_service, docx_service
from app.services.errors import PermissionDenied, ValidationFailed


@pytest.fixture()
def demanda(db, reader, area):
    cr = change_request_service.create(
        db, reader, title="Reprovação sem valor de pedido", area_id=area.id,
        description_html='<p>Regra <strong>soberana</strong> no início do fluxo.</p>',
    )
    db.commit()
    return cr


def test_create_sanitiza_html(db, reader, area):
    cr = change_request_service.create(
        db, reader, title="XSS", area_id=area.id,
        description_html='<p>ok</p><script>alert(1)</script>',
    )
    assert "<script" not in cr.description_html
    assert "ok" in cr.description_html


def test_update_pelo_solicitante_e_historico(db, reader, demanda):
    change_request_service.update(
        db, reader, demanda.id,
        title="Reprovação sem valor de pedido (F4)",
        description_html="<h2>Escopo</h2><p>Somente origem F4.</p>",
        priority="high",
    )
    db.commit()
    assert demanda.title.endswith("(F4)")
    assert demanda.priority == "high"
    history = change_request_service.update_history(db, demanda.id)
    actions = [h.action for h in history]
    assert "change_request.updated" in actions
    assert "change_request.created" in actions


def test_update_por_outro_leitor_negado(db, reader, demanda, admin):
    from app.services import user_service

    outro = user_service.create_user(
        db, admin, username="outro", email="outro@example.com", display_name="Outro",
        role="reader", password="senha-forte-123", must_change_password=False,
    )
    db.commit()
    with pytest.raises(PermissionDenied):
        change_request_service.update(
            db, outro, demanda.id, title="x", description_html="", priority="low",
        )


def test_update_demanda_encerrada_falha(db, reader, approver, demanda):
    change_request_service.reject(db, approver, demanda.id, "duplicada")
    db.commit()
    with pytest.raises(ValidationFailed):
        change_request_service.update(
            db, reader, demanda.id, title="x", description_html="", priority="low",
        )


def test_anexo_na_demanda(db, reader, demanda):
    attachment = attachment_service.upload_for_change_request(
        db, reader, demanda.id,
        filename="matriz.csv", content=b"a;b\n1;2\n", content_type="text/csv",
    )
    db.commit()
    assert attachment.change_request_id == demanda.id
    assert attachment.version_id is None
    got, content = attachment_service.get_content(db, reader, attachment.id)
    assert content == b"a;b\n1;2\n"


def test_anexo_extensao_proibida(db, reader, demanda):
    with pytest.raises(ValidationFailed):
        attachment_service.upload_for_change_request(
            db, reader, demanda.id, filename="virus.exe", content=b"x",
        )


def test_export_docx_com_cineminha(db, reader, author, area, demanda):
    from app.services import cinema_service

    var_g = cinema_service.create_variable(db, author, name="GRUPO", label="Grupo", domain="G7")
    var_s = cinema_service.create_variable(
        db, author, name="SCORE", label="Score", domain="R01, R02"
    )
    cinema = cinema_service.create_cinema(
        db, author, name="Cortes G7", cinema_type="offer",
        row_variable_id=var_g.id, col_variable_id=var_s.id,
    )
    cinema_service.create_manual_version(db, author, cinema.id, cells={"G7|R01": 8})
    instance = cinema_service.add_instance(db, author, demanda.id, cinema.id)
    cinema_service.update_instance_cells(db, author, instance.id, {"G7|R01": 5, "G7|R02": 3})
    db.commit()

    content = docx_service.export_change_request_docx(db, demanda)
    assert content[:2] == b"PK"  # zip/docx

    import docx as docx_lib

    document = docx_lib.Document(io.BytesIO(content))
    text = "\n".join(p.text for p in document.paragraphs)
    assert demanda.code in text
    assert "Cortes G7" in text
    # matriz vira tabela nativa (metadados + registro + matriz)
    cells = [c.text for t in document.tables for r in t.rows for c in r.cells]
    assert "5" in cells and "3" in cells
