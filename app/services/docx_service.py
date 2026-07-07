"""Exportação de demanda para Word (.docx) via python-docx.

Gera o documento que hoje é montado à mão no Word: cabeçalho da demanda,
registro de alterações (da trilha de auditoria), corpo formatado (HTML do
editor WYSIWYG convertido num subconjunto de estruturas do Word) e cada
cineminha como tabela nativa colorida — com destaque nas caselas alteradas
em relação à versão de origem da biblioteca.

python-docx é dependência pura de Python (sem libs de sistema); o import é
tardio para o restante do sistema operar mesmo sem ela instalada.
"""

from __future__ import annotations

import base64
import io
import re
from html.parser import HTMLParser

from sqlalchemy.orm import Session

from app.models import ChangeRequest, CinemaType
from app.services import change_request_service, cinema_service
from app.services.errors import ValidationFailed

# ── Cores das matrizes (hex sem '#', formato python-docx) ────────────────────
_GREEN = "63BE7B"
_YELLOW = "FFEB84"
_RED = "F8696B"
_GRAY = "F1F5F9"
_HEADER_BG = "4F46E5"
_CHANGED_BORDER = "1D4ED8"


def _require_docx():
    try:
        import docx  # noqa: F401
    except ImportError as exc:  # pragma: no cover - depende do ambiente
        raise ValidationFailed(
            "exportação Word indisponível: instale a dependência python-docx"
        ) from exc


# ── HTML (editor) → parágrafos/tabelas do Word ──────────────────────────────


class _DocxHTMLWriter(HTMLParser):
    """Converte o subconjunto sanitizado de HTML em conteúdo python-docx."""

    _HEADINGS = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}

    def __init__(self, document) -> None:
        super().__init__()
        self.doc = document
        self.paragraph = None
        self.bold = 0
        self.italic = 0
        self.underline = 0
        self.strike = 0
        self.list_stack: list[str] = []  # 'ol' | 'ul'
        self.in_table = False
        self.table_rows: list[list[str]] = []
        self.current_row: list[str] | None = None
        self.current_cell: list[str] | None = None
        self.pre_depth = 0

    # -- helpers -------------------------------------------------------------

    def _ensure_paragraph(self, style: str | None = None):
        if self.paragraph is None:
            self.paragraph = self.doc.add_paragraph(style=style)
        return self.paragraph

    def _close_paragraph(self) -> None:
        self.paragraph = None

    def _add_text(self, text: str) -> None:
        if not text:
            return
        if self.current_cell is not None:
            self.current_cell.append(text)
            return
        paragraph = self._ensure_paragraph()
        run = paragraph.add_run(text)
        run.bold = self.bold > 0 or None
        run.italic = self.italic > 0 or None
        run.underline = self.underline > 0 or None
        if self.strike:
            run.font.strike = True

    # -- eventos -------------------------------------------------------------

    def handle_starttag(self, tag: str, attrs) -> None:
        attrs_dict = dict(attrs)
        if tag in self._HEADINGS:
            self._close_paragraph()
            self.paragraph = self.doc.add_heading("", level=self._HEADINGS[tag])
        elif tag == "p":
            self._close_paragraph()
        elif tag == "br":
            if self.paragraph is not None:
                self.paragraph.add_run().add_break()
        elif tag in ("strong", "b"):
            self.bold += 1
        elif tag in ("em", "i"):
            self.italic += 1
        elif tag == "u":
            self.underline += 1
        elif tag == "s":
            self.strike += 1
        elif tag in ("ol", "ul"):
            self.list_stack.append(tag)
            self._close_paragraph()
        elif tag == "li":
            self._close_paragraph()
            # Quill marca <li data-list="bullet|ordered"> dentro de <ol>
            kind = attrs_dict.get("data-list") or (
                "ordered" if (self.list_stack and self.list_stack[-1] == "ol") else "bullet"
            )
            style = "List Number" if kind == "ordered" else "List Bullet"
            self.paragraph = self.doc.add_paragraph(style=style)
        elif tag == "blockquote":
            self._close_paragraph()
            self.paragraph = self.doc.add_paragraph(style="Intense Quote")
        elif tag == "pre":
            self.pre_depth += 1
            self._close_paragraph()
            self.paragraph = self.doc.add_paragraph()
        elif tag == "table":
            self.in_table = True
            self.table_rows = []
        elif tag == "tr" and self.in_table:
            self.current_row = []
        elif tag in ("td", "th") and self.in_table:
            self.current_cell = []
        elif tag == "img":
            self._add_image(attrs_dict.get("src") or "")
        elif tag == "hr":
            self._close_paragraph()
            self.doc.add_paragraph("—" * 30)

    def handle_endtag(self, tag: str) -> None:
        if tag in self._HEADINGS or tag in ("p", "li", "blockquote"):
            self._close_paragraph()
        elif tag in ("strong", "b"):
            self.bold = max(0, self.bold - 1)
        elif tag in ("em", "i"):
            self.italic = max(0, self.italic - 1)
        elif tag == "u":
            self.underline = max(0, self.underline - 1)
        elif tag == "s":
            self.strike = max(0, self.strike - 1)
        elif tag in ("ol", "ul"):
            if self.list_stack:
                self.list_stack.pop()
            self._close_paragraph()
        elif tag == "pre":
            self.pre_depth = max(0, self.pre_depth - 1)
            self._close_paragraph()
        elif tag in ("td", "th") and self.current_row is not None:
            self.current_row.append("".join(self.current_cell or []).strip())
            self.current_cell = None
        elif tag == "tr" and self.in_table and self.current_row is not None:
            self.table_rows.append(self.current_row)
            self.current_row = None
        elif tag == "table" and self.in_table:
            self._flush_table()
            self.in_table = False

    def handle_data(self, data: str) -> None:
        if self.pre_depth:
            self._add_text(data)
        else:
            self._add_text(re.sub(r"\s+", " ", data))

    # -- estruturas ----------------------------------------------------------

    def _flush_table(self) -> None:
        rows = [r for r in self.table_rows if r]
        if not rows:
            return
        cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            for j in range(cols):
                table.rows[i].cells[j].text = row[j] if j < len(row) else ""
        self._close_paragraph()

    def _add_image(self, src: str) -> None:
        match = re.match(r"^data:image/(png|jpe?g|gif);base64,(.+)$", src, re.DOTALL)
        if not match:
            return  # só imagens embutidas (data URL) — links externos são ignorados
        try:
            payload = base64.b64decode(match.group(2))
        except (ValueError, TypeError):
            return
        from docx.shared import Inches

        self._close_paragraph()
        try:
            self.doc.add_picture(io.BytesIO(payload), width=Inches(6.0))
        except Exception:  # imagem corrompida não derruba o export
            self.doc.add_paragraph("[imagem não suportada]")


# ── Matriz do cineminha como tabela colorida ─────────────────────────────────


def _shade(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def _cell_color(cinema_type: str, value, max_value: float) -> str:
    """Mesma semântica de cor da tela: elegível verde/vermelho; oferta em escala."""
    number = float(value)
    if cinema_type == CinemaType.ELIGIBILITY:
        return _GREEN if number >= 1 else _RED
    if number <= 0:
        return _RED
    if max_value <= 0:
        return _GRAY
    ratio = number / max_value
    if ratio >= 0.66:
        return _GREEN
    if ratio >= 0.33:
        return _YELLOW
    return _RED


def _mark_changed(cell) -> None:
    """Borda grossa azul na casela alterada vs. a origem da biblioteca."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "18")
        el.set(qn("w:color"), _CHANGED_BORDER)
        borders.append(el)
    cell._tc.get_or_add_tcPr().append(borders)


def _add_matrix_table(document, view: dict, row_label: str, col_label: str) -> None:
    from docx.shared import Pt

    rows = view["row_domain"]
    cols = view["col_domain"]
    table = document.add_table(rows=len(rows) + 1, cols=len(cols) + 1)
    table.style = "Table Grid"

    header = table.rows[0]
    header.cells[0].text = f"{row_label} \\ {col_label}"
    _shade(header.cells[0], _HEADER_BG)
    for j, col in enumerate(cols):
        cell = header.cells[j + 1]
        cell.text = str(col)
        _shade(cell, _HEADER_BG)

    eligibility = view["cinema_type"] == CinemaType.ELIGIBILITY
    for i, row_value in enumerate(rows):
        row = table.rows[i + 1]
        row.cells[0].text = str(row_value)
        _shade(row.cells[0], _GRAY)
        for j in range(len(cols)):
            value = view["grid"][i][j]
            cell = row.cells[j + 1]
            if eligibility:
                cell.text = "✓" if float(value) >= 1 else "✗"
            else:
                cell.text = f"{value:g}" if isinstance(value, float) else str(value)
            _shade(cell, _cell_color(view["cinema_type"], value, view["max_value"]))
            if view["changed"][i][j]:
                _mark_changed(cell)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1  # centralizado
                for run in paragraph.runs:
                    run.font.size = Pt(8)


# ── Documento completo da demanda ────────────────────────────────────────────

_PRIORITY_LABELS = {
    "low": "Baixa",
    "medium": "Média",
    "high": "Alta",
    "regulatory": "Regulatória",
}

_STATUS_LABELS = {
    "open": "Aberta",
    "in_progress": "Em andamento",
    "done": "Concluída",
    "rejected": "Rejeitada",
}

_AUDIT_ACTION_LABELS = {
    "change_request.created": "Demanda registrada",
    "change_request.updated": "Demanda editada",
    "change_request.started": "Marcada em andamento",
    "change_request.done": "Concluída pela vigência",
    "change_request.rejected": "Rejeitada",
}


def export_change_request_docx(db: Session, change_request: ChangeRequest) -> bytes:
    """Monta o .docx completo da demanda e retorna os bytes."""
    _require_docx()
    from docx import Document

    document = Document()
    document.add_heading("Solicitação de Mudança em Políticas", level=0)
    document.add_heading(f"{change_request.code} — {change_request.title}", level=1)

    # cabeçalho de metadados
    meta = document.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    meta_rows = [
        ("Solicitante", change_request.requester.display_name if change_request.requester else ""),
        ("Área", change_request.area.name if change_request.area else ""),
        ("Prioridade", _PRIORITY_LABELS.get(change_request.priority, change_request.priority)),
        ("Status", _STATUS_LABELS.get(change_request.status, change_request.status)),
        ("Aberta em", change_request.created_at.strftime("%d/%m/%Y")),
    ]
    for i, (label, value) in enumerate(meta_rows):
        meta.rows[i].cells[0].text = label
        meta.rows[i].cells[1].text = value or ""
        _shade(meta.rows[i].cells[0], _GRAY)

    # registro de alterações (trilha de auditoria)
    history = change_request_service.update_history(db, change_request.id)
    if history:
        document.add_heading("Registro de alterações", level=2)
        table = document.add_table(rows=len(history) + 1, cols=3)
        table.style = "Table Grid"
        for j, head in enumerate(("Data", "Ação", "Autor")):
            table.rows[0].cells[j].text = head
            _shade(table.rows[0].cells[j], _GRAY)
        for i, entry in enumerate(history):
            table.rows[i + 1].cells[0].text = entry.created_at.strftime("%d/%m/%Y %H:%M")
            table.rows[i + 1].cells[1].text = _AUDIT_ACTION_LABELS.get(
                entry.action, entry.action
            )
            table.rows[i + 1].cells[2].text = entry.actor.display_name if entry.actor else "sistema"

    # corpo (HTML do editor)
    if change_request.description_html or change_request.description_md:
        document.add_heading("Descrição", level=2)
        if change_request.description_html:
            writer = _DocxHTMLWriter(document)
            writer.feed(change_request.description_html)
        else:
            for line in change_request.description_md.splitlines():
                document.add_paragraph(line)

    # cineminhas da demanda
    instances = cinema_service.list_instances(db, change_request.id)
    if instances:
        document.add_heading("Cineminhas (matrizes de política)", level=2)
        for instance in instances:
            view = cinema_service.instance_view(instance)
            cinema = instance.cinema
            source = instance.source_version
            document.add_heading(cinema.name, level=3)
            type_label = (
                "Elegibilidade" if cinema.cinema_type == CinemaType.ELIGIBILITY else "Oferta"
            )
            origin = (
                f"origem: v{source.version_number} da biblioteca"
                if source
                else "origem: biblioteca sem versão vigente (matriz nova)"
            )
            document.add_paragraph(f"Tipo: {type_label} · {origin}")
            if view["changed_count"]:
                document.add_paragraph(
                    f"{view['changed_count']} casela(s) alterada(s) em relação à origem "
                    "(destacadas com borda azul)."
                )
            _add_matrix_table(
                document,
                view,
                cinema.row_variable.label or cinema.row_variable.name,
                cinema.col_variable.label or cinema.col_variable.name,
            )
            if instance.notes:
                document.add_paragraph(f"Observações: {instance.notes}")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
